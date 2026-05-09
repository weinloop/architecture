#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Markdown → DOCX：Pandoc 转换后对表格加实线、整体紧凑版面。"""

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent

MD_TARGETS = (
    ROOT / "101综合知识汇总.md",
    ROOT / "101案例汇总.md",
)


def tbl_set_borders(table, sz_twips="6", color="000000"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for child in list(tblPr.findall(qn("w:tblBorders"))):
        tblPr.remove(child)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz_twips)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def tbl_compact_cellmargins(table):
    """减小表元内边距（单位：twips，1 pt = 20 twips）"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for child in list(tblPr.findall(qn("w:tblCellMar"))):
        tblPr.remove(child)
    mar = OxmlElement("w:tblCellMar")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), "40")  # 约 2 pt
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)


def iter_every_paragraph(doc):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def tighten_document(doc, body_pt=10.5):
    """段前段后压低、正文约五号、单倍行距；页边距略收。"""
    for sec in doc.sections:
        sec.left_margin = Cm(1.5)
        sec.right_margin = Cm(1.5)
        sec.top_margin = Cm(1.6)
        sec.bottom_margin = Cm(1.6)

    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        try:
            pf = style.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(2)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except (AttributeError, TypeError):
            pass
        if style.name == "Normal":
            try:
                style.font.size = Pt(body_pt)
                style.font.name = "宋体"
            except (AttributeError, TypeError):
                pass

    for name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title"):
        try:
            s = doc.styles[name]
            s.paragraph_format.space_before = Pt(4)
            s.paragraph_format.space_after = Pt(2)
            s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except KeyError:
            pass

    for p in iter_every_paragraph(doc):
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        try:
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except TypeError:
            pass
        for r in p.runs:
            try:
                if r.font.size is None:
                    r.font.size = Pt(body_pt)
                if not r.font.name:
                    r.font.name = "宋体"
            except (AttributeError, TypeError):
                pass


def postprocess_docx(out: Path, body_pt=10.5) -> None:
    """对已有 docx：表线、紧凑排版。"""
    doc = Document(str(out))
    for tbl in doc.tables:
        tbl_set_borders(tbl)
        tbl_compact_cellmargins(tbl)
    tighten_document(doc, body_pt=body_pt)
    doc.save(str(out))


def md_to_docx(md_path: Path) -> Path:
    out = md_path.with_suffix(".docx")
    subprocess.run(
        [
            "pandoc",
            "-f",
            "markdown-yaml_metadata_block+pipe_tables+grid_tables+fenced_divs",
            "-t",
            "docx",
            "--standalone",
            "--wrap=none",
            "-o",
            str(out),
            str(md_path),
        ],
        check=True,
    )
    postprocess_docx(out)
    return out


def main(argv: list[str]) -> None:
    paths = MD_TARGETS if not argv else [Path(a) for a in argv]
    for md in paths:
        if not md.is_file():
            print("跳过（不存在）：", md, file=sys.stderr)
            continue
        out = md_to_docx(md)
        print(out)


if __name__ == "__main__":
    main(sys.argv[1:])
