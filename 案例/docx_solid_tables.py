#!/usr/bin/env python3
"""为 Word 表格设置实线外框与内部分隔线。"""
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
DOC = Path(sys.argv[1]) if len(sys.argv) > 1 else HERE / "案例汇总.docx"


def tbl_set_borders(table, sz_twips="8", color="000000"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for child in list(tblPr.findall(qn("w:tblBorders"))):
        tblPr.remove(child)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz_twips)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def main():
    doc = Document(DOC)
    for tbl in doc.tables:
        tbl_set_borders(tbl)
    for sec in doc.sections:
        for part in (sec.header, sec.footer):
            if part is not None:
                for tbl in part.tables:
                    tbl_set_borders(tbl)
    doc.save(DOC)
    print("Applied solid borders:", DOC)


if __name__ == "__main__":
    main()
